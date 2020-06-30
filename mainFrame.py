import tkinter as tk
from tkinter import ttk
import sqlite3
import datetime
import random
from docx import Document


class Main(tk.Frame):
    def __init__(self, root):
        super().__init__(root)
        self.init_main()
        self.db = db
        self.view_records()

    def init_main(self):
        toolbar = tk.Frame(bg='#d7d8e0', bd=2)
        toolbar.pack(side=tk.TOP, fill=tk.X)

        self.add_img = tk.PhotoImage(file='./img/add.png')
        btn_open_dialog = tk.Button(toolbar, text='Добавить', command=self.open_dialog, bg='#d7d8e0', bd=0,
                                    compound=tk.TOP, image=self.add_img)
        btn_open_dialog.pack(side=tk.LEFT)

        self.update_img = tk.PhotoImage(file='./img/refactor.png')
        btn_edit_dialog = tk.Button(toolbar, text='Редактировать', bg='#d7d8e0', bd=0, image=self.update_img,
                                    compound=tk.TOP, command=self.open_update_dialog)
        btn_edit_dialog.pack(side=tk.LEFT)

        self.delete_img = tk.PhotoImage(file='./img/delete.png')
        btn_delete = tk.Button(toolbar, text='Удалить', bg='#d7d8e0', bd=0, image=self.delete_img,
                               compound=tk.TOP, command=self.delete_records)
        btn_delete.pack(side=tk.LEFT)

        self.search_img = tk.PhotoImage(file='./img/search.png')
        btn_search = tk.Button(toolbar, text='Поиск', bg='#d7d8e0', bd=0, image=self.search_img,
                               compound=tk.TOP, command=self.open_search_dialog)
        btn_search.pack(side=tk.LEFT)

        self.refresh_img = tk.PhotoImage(file='./img/refresh.png')
        btn_refresh = tk.Button(toolbar, text='Обновить', bg='#d7d8e0', bd=0, image=self.refresh_img,
                                compound=tk.TOP, command=self.view_records)
        btn_refresh.pack(side=tk.LEFT)

        self.buy_img = tk.PhotoImage(file='./img/buy.png')
        btn_buy = tk.Button(toolbar, text='Оформить', bg='#d7d8e0', bd=0, image=self.buy_img,
                            compound=tk.TOP, command=self.open_buy)
        btn_buy.pack(side=tk.RIGHT)

        self.orders_img = tk.PhotoImage(file='./img/orders.png')
        btn_orders = tk.Button(toolbar, text='Покупки', bg='#d7d8e0', bd=0, image=self.orders_img,
                               compound=tk.TOP, command=self.open_orders)
        btn_orders.pack(side=tk.RIGHT)

        self.tree = ttk.Treeview(self, columns=('ID', 'country', 'city', 'name_tour', 'date', 'quantity', 'price'),
                                 height=50, show='headings')

        self.tree.column('ID', width=30, anchor=tk.CENTER)
        self.tree.column('country', width=150, anchor=tk.CENTER)
        self.tree.column('city', width=150, anchor=tk.CENTER)
        self.tree.column('name_tour', width=150, anchor=tk.CENTER)
        self.tree.column('date', width=120, anchor=tk.CENTER)
        self.tree.column('quantity', width=120, anchor=tk.CENTER)
        self.tree.column('price', width=140, anchor=tk.CENTER)

        self.tree.heading('ID', text='ID')
        self.tree.heading('country', text='Страна')
        self.tree.heading('city', text='Город')
        self.tree.heading('name_tour', text='Название тура')
        self.tree.heading('date', text='Дата вылета')
        self.tree.heading('quantity', text='Кол-во билетов')
        self.tree.heading('price', text='Цена билета')

        self.tree.pack()

    def records(self, country, city, name_tour, date, quantity, price):
        self.db.insert_data(country, city, name_tour, date, quantity, price)
        self.view_records()

    def buy(self, country, city, name_tour, date, quantity, price):
        now = datetime.datetime.now()
        now = now.strftime("%d-%m-%Y %H:%M")
        rnd = ''
        summa = int(quantity) * int(price)

        for x in range(4):
            rnd = rnd + random.choice(list('123456789qwertyuiopasdfghjklzxcvbnmQWERTYUIOPASDFGHJKLZXCVBNM'))

        self.db.c.execute(
            '''UPDATE tours SET country =?, city=?, name_tour=?, date=?, quantity=quantity - ?, price=? WHERE ID=?''',
            (country, city, name_tour, date, quantity, price,
             self.tree.set(self.tree.selection()[0], '#1')))
        self.db.c.execute(
            '''INSERT INTO orders(name_order, q_order, summa, date_order, date_out) VALUES (?, ?, ?, ?, ?)''',
            (rnd, quantity, summa, now, date))

        document = Document('data/order.docx')
        document.add_paragraph('{} {}'.format('Код договора:', rnd))
        document.add_paragraph('{} {}'.format('Дата:', now))
        document.add_paragraph('{} {}'.format('Кол-во билетов:', quantity))
        document.add_paragraph('{} {}'.format('Тур:', name_tour))
        document.add_paragraph('{} {}'.format('Сумма:', summa))
        document.add_paragraph('{} {}'.format('Вылет:', date))
        new_str = '{}_{}{}'.format('data/order', rnd, '.docx')
        document.save(new_str)
        self.db.conn.commit()

        self.db.conn.commit()
        self.view_records()

    def update_record(self, country, city, name_tour, date, quantity, price):
        self.db.c.execute(
            '''UPDATE tours SET country =?, city=?, name_tour=?,date=?, quantity=?, price=? WHERE ID=?''',
            (country, city, name_tour, date, quantity, price,
             self.tree.set(self.tree.selection()[0], '#1')))
        self.db.conn.commit()
        self.view_records()

    def view_records(self):
        self.db.c.execute('''SELECT * FROM tours''')
        [self.tree.delete(i) for i in self.tree.get_children()]
        [self.tree.insert('', 'end', values=row) for row in self.db.c.fetchall()]
        self.tree.config(height=len(self.tree.get_children()))

    def delete_records(self):
        for selection_item in self.tree.selection():
            self.db.c.execute('''DELETE FROM tours WHERE id=?''', (self.tree.set(selection_item, '#1'),))
        self.db.conn.commit()
        self.view_records()

    def search_records(self, country):
        country = ('%' + country + '%',)
        self.db.c.execute('''SELECT * FROM tours WHERE country LIKE ?''', country)
        [self.tree.delete(i) for i in self.tree.get_children()]
        [self.tree.insert('', 'end', values=row) for row in self.db.c.fetchall()]

    def open_dialog(self):
        Child()

    def open_update_dialog(self):
        Update()

    def open_search_dialog(self):
        Search()

    def open_buy(self):
        Buy()

    def open_orders(self):
        Orders()


class Child(tk.Toplevel):
    def __init__(self):
        super().__init__(root)
        self.init_child()
        self.view = app

    def init_child(self):
        self.title('Добавить тур')
        self.geometry('400x300+400+300')
        self.resizable(False, False)

        label_country = tk.Label(self, text='Страна:')
        label_country.place(x=50, y=50)
        label_city = tk.Label(self, text='Город:')
        label_city.place(x=50, y=80)
        label_name_tour = tk.Label(self, text='Название тура:')
        label_name_tour.place(x=50, y=110)
        label_date = tk.Label(self, text='Дата вылета:')
        label_date.place(x=50, y=140)
        label_quantity = tk.Label(self, text='Кол-во билетов:')
        label_quantity.place(x=50, y=170)
        label_price = tk.Label(self, text='Цена билета:')
        label_price.place(x=50, y=200)

        self.entry_country = ttk.Entry(self)
        self.entry_country.place(x=200, y=50)

        self.entry_city = ttk.Entry(self)
        self.entry_city.place(x=200, y=80)

        self.entry_name_tour = ttk.Entry(self, )
        self.entry_name_tour.place(x=200, y=110)

        self.entry_date = ttk.Entry(self)
        self.entry_date.place(x=200, y=140)

        self.entry_quantity = ttk.Entry(self)
        self.entry_quantity.place(x=200, y=170)

        self.entry_price = ttk.Entry(self)
        self.entry_price.place(x=200, y=200)

        btn_cancel = ttk.Button(self, text='Закрыть', command=self.destroy)
        btn_cancel.place(x=300, y=240)

        self.btn_ok = ttk.Button(self, text='Добавить')
        self.btn_ok.place(x=220, y=240)
        self.btn_ok.bind('<Button-1>', lambda event: self.view.records(self.entry_country.get(),
                                                                       self.entry_city.get(),
                                                                       self.entry_name_tour.get(),
                                                                       self.entry_date.get(),
                                                                       self.entry_quantity.get(),
                                                                       self.entry_price.get()))

        self.grab_set()
        self.focus_set()


class Orders(tk.Toplevel):
    def __init__(self):
        super().__init__()
        self.init_orders()
        self.view = app
        self.db = db
        self.view_records_orders()

    def init_orders(self):
        self.title("Покупки")
        self.tree2 = ttk.Treeview(self, columns=('ID', 'name_order', 'q_order', 'summa', 'date_order', 'date_out'),
                                  height=50, show='headings')

        self.tree2.column('ID', width=30, anchor=tk.CENTER)
        self.tree2.column('name_order', width=150, anchor=tk.CENTER)
        self.tree2.column('q_order', width=150, anchor=tk.CENTER)
        self.tree2.column('summa', width=150, anchor=tk.CENTER)
        self.tree2.column('date_order', width=150, anchor=tk.CENTER)
        self.tree2.column('date_out', width=150, anchor=tk.CENTER)

        self.tree2.heading('ID', text='ID')
        self.tree2.heading('name_order', text='Код договора')
        self.tree2.heading('q_order', text='Кол-во билетов')
        self.tree2.heading('summa', text='Сумма')
        self.tree2.heading('date_order', text='Дата оформления')
        self.tree2.heading('date_out', text='Дата вылета')

        self.tree2.pack()

    def records_orders(self, name_order, q_order, summa, date_order, date_out):
        self.db.insert_data_orders(name_order, q_order, summa, date_order, date_out)
        self.view_records_orders()

    def view_records_orders(self):
        self.db.c.execute('''SELECT * FROM orders''')
        [self.tree2.delete(i) for i in self.tree2.get_children()]
        [self.tree2.insert('', 'end', values=row) for row in self.db.c.fetchall()]
        self.tree2.config(height=len(self.tree2.get_children()))

    def select_pass(self):
        self.db.c.execute('''SELECT pass FROM pass''')


class Buy(Child):
    def __init__(self):
        super().__init__()
        self.init_edit()
        self.view = app
        self.db = db
        self.default_data()

    def init_edit(self):
        self.title('Оформление покупки')
        btn_edit = ttk.Button(self, text='Оформить заказ')
        btn_edit.place(x=205, y=240)
        btn_edit.bind('<Button-1>', lambda event: self.view.buy(self.entry_country.get(),
                                                                self.entry_city.get(),
                                                                self.entry_name_tour.get(),
                                                                self.entry_date.get(),
                                                                self.entry_quantity.get(),
                                                                self.entry_price.get()))

    def default_data(self):
        self.db.c.execute('''SELECT * FROM tours WHERE id=?''',
                          (self.view.tree.set(self.view.tree.selection()[0], '#1')))
        row = self.db.c.fetchone()
        self.entry_country.insert(0, row[1])
        self.entry_city.insert(0, row[2])
        self.entry_name_tour.insert(0, row[3])
        self.entry_date.insert(0, row[4])
        self.entry_quantity
        self.entry_price.insert(0, row[6])


class Update(Child):
    def __init__(self):
        super().__init__()
        self.init_edit()
        self.view = app
        self.db = db
        self.default_data()

    def init_edit(self):
        self.title('Редактировать')
        btn_edit = ttk.Button(self, text='Редактировать')
        btn_edit.place(x=205, y=240)
        btn_edit.bind('<Button-1>', lambda event: self.view.update_record(self.entry_country.get(),
                                                                          self.entry_city.get(),
                                                                          self.entry_name_tour.get(),
                                                                          self.entry_date.get(),
                                                                          self.entry_quantity.get(),
                                                                          self.entry_price.get()))

        self.btn_ok.destroy()

    def default_data(self):
        self.db.c.execute('''SELECT * FROM tours WHERE id=?''',
                          (self.view.tree.set(self.view.tree.selection()[0], '#1')))
        row = self.db.c.fetchone()
        self.entry_country.insert(0, row[1])
        self.entry_city.insert(0, row[2])
        self.entry_name_tour.insert(0, row[3])
        self.entry_date.insert(0, row[4])
        self.entry_quantity.insert(0, row[5])
        self.entry_price.insert(0, row[6])


class Search(tk.Toplevel):
    def __init__(self):
        super().__init__()
        self.init_search()
        self.view = app

    def init_search(self):
        self.title('Поиск')
        self.geometry('300x100+400+300')
        self.resizable(False, False)

        label_search = tk.Label(self, text='Поиск')
        label_search.place(x=50, y=20)

        self.entry_search = ttk.Entry(self)
        self.entry_search.place(x=105, y=20, width=150)

        btn_cancel = ttk.Button(self, text='Закрыть', command=self.destroy)
        btn_cancel.place(x=185, y=50)

        btn_search = ttk.Button(self, text='Поиск')
        btn_search.place(x=105, y=50)
        btn_search.bind('<Button-1>', lambda event: self.view.search_records(self.entry_search.get()))
        btn_search.bind('<Button-1>', lambda event: self.destroy(), add='+')


class DB:

    def __init__(self):
        self.conn = sqlite3.connect('tours.db')
        self.c = self.conn.cursor()

        self.c.execute(
            '''CREATE TABLE IF NOT EXISTS tours
             (id integer primary key,
              country text,
              city text, 
              name_tour text,
              date date,     
              quantity integer,
              price integer,
              CONSTRAINT country_unique UNIQUE (country, city))''')
        self.conn.commit()

        self.c.execute('''CREATE TABLE IF NOT EXISTS orders
                       (id integer primary key,
                        name_order text,
                        q_order integer,
                        summa integer,
                        date_order date,
                        date_out date)''')
        self.conn.commit()
        self.c.execute('''CREATE TABLE IF NOT EXISTS pass
                               (id integer primary key,
                                pass text)''')
        self.conn.commit()

    def insert_data(self, country, city, name_tour, date, quantity, price):
        self.c.execute(
            '''INSERT INTO tours(country, city, name_tour, date, quantity, price) VALUES (?, ?, ?, ?, ?, ?)''',
            (country, city, name_tour, date, quantity, price))
        self.conn.commit()

    def insert_data_orders(self, name_order, q_order, summa, date_order, date_out):
        self.c.execute(
            '''INSERT INTO orders(name_order, q_order, summa, date_order,date_out) VALUES (?, ?, ?, ?, ?)''',
            (name_order, q_order, summa, date_order, date_out))
        self.conn.commit()


if __name__ == "__main__":
    root = tk.Tk()
    db = DB()
    app = Main(root)
    app.pack()
    root.title("v_1")
    root.geometry("750x500+300+200")
    root.resizable(True, True)
    root.mainloop()
